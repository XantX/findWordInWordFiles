# Buscar todos los documentos

import os
import docx
import shutil


def findInDocument(fileName, target):
    doc = docx.Document(fileName)
    for paragraph in doc.paragraphs:
        index = paragraph.text.find(target)
        if index >= 0:
            return fileName
    return 0


folder = input("Escriba la ruta donde estan almacenados los archivos word: ")
input = input("Escriba el codigo que desea buscar: ")

for path, dirs, files in os.walk(folder):
    for f in files:
        fileName = os.path.join(path, f)
        print("Buscando en:", fileName)
        result = findInDocument(fileName, input)
        if result != 0:
            print("Encontrado en:", fileName)
            shutil.copy(fileName, "./result")
            quit()
